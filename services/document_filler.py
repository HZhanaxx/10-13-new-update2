"""
Document AutoFiller Service
===========================

Fills legal document templates with data from questionnaire answers.
Integrates with:
- Database (document_templates, documents, document_data)
- OCR results parsing
- LangGraph workflow

Based on the input.py filler logic, adapted for web service use.
"""

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from collections import defaultdict
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import json
import re
import os
import uuid
import tempfile
import shutil
from datetime import datetime

from sqlalchemy.orm import Session
from pydantic import BaseModel, Field


# ============================================================================
# Pydantic Models
# ============================================================================

class TemplateField(BaseModel):
    """Definition of a field required by a template"""
    field_name: str
    field_type: str = "text"  # text, date, number, array, client_info
    required: bool = True
    description: Optional[str] = None
    example: Optional[str] = None


class FillerRequest(BaseModel):
    """Request to fill a template"""
    template_code: str
    data: Dict[str, Any]
    apply_fangsong: bool = True
    fangsong_font: str = "仿宋_GB2312"
    output_filename: Optional[str] = None


class FillerResult(BaseModel):
    """Result of template filling"""
    success: bool
    output_path: Optional[str] = None
    output_filename: Optional[str] = None
    filled_fields: int = 0
    error: Optional[str] = None


# ============================================================================
# JSON TRANSFORMATION FUNCTIONS (from input.py)
# ============================================================================

def has_oppoclient_data(input_data: dict) -> bool:
    """Check if there are any OppoClient fields in the input data."""
    for key in input_data.keys():
        if key.startswith("OppoClientName"):
            suffix = key[len("OppoClientName"):]
            if suffix and suffix.isdigit():
                return True
    return False


def extract_numbered_fields(data: dict, prefix: str) -> dict:
    """
    Extract all fields with a given prefix and organize by number.
    Returns a dict where keys are numbers and values are dicts of field_suffix: value
    """
    clients = defaultdict(dict)
    
    for key, value in data.items():
        if key.startswith(prefix):
            for i in range(len(key)-1, -1, -1):
                if key[i].isdigit():
                    num_start = i
                    while num_start > 0 and key[num_start-1].isdigit():
                        num_start -= 1
                    
                    number = key[num_start:]
                    field_name = key[len(prefix):num_start]
                    
                    if number.isdigit():
                        clients[number][field_name] = value
                    break
            else:
                field_name = key[len(prefix):]
                if field_name:
                    clients["0"][field_name] = value
    
    return dict(clients)


def format_oriclient(client_data: dict, form: str = "异议", is_opposing: bool = False) -> str:
    """Format original client information."""
    status = client_data.get("Status", form)
    name = client_data.get("Name", "")
    gender = client_data.get("Gender", "")
    dob = client_data.get("DOB", "")
    race = client_data.get("Race", "")
    workdet = client_data.get("Workdet", "")
    address = client_data.get("Address", "")
    tele = client_data.get("Tele", "")
    
    if is_opposing:
        form_label = f"被{form}人"
    else:
        form_label = f"{form}人"
    
    return f"{form_label}({status}):{name},{gender},{dob}出生,{race},系{workdet},住址{address},联系方式{tele}。"


def format_realclient(client_data: dict) -> str:
    """Format real client (代理人) information."""
    rank = client_data.get("Rank", "指定")
    name = client_data.get("Name", "")
    gender = client_data.get("Gender", "")
    dob = client_data.get("DOB", "")
    race = client_data.get("Race", "")
    address = client_data.get("Address", "")
    tele = client_data.get("Tele", "")
    relationship = client_data.get("Relationship", "")
    
    return f"{rank}代理人:{name},{gender},{dob}出生,{race},住址{address},联系方式{tele},系异议人之{relationship}。"


def format_agent(agent_data: dict) -> str:
    """Format agent information."""
    name = agent_data.get("Name", "")
    workdet = agent_data.get("Workdet", "")
    
    return f"委托诉讼代理人:{name},{workdet}。"


def summarize_client_names(input_data: dict) -> Tuple[List[str], List[str]]:
    """
    Summarize all numbered client names into single fields.
    OriClientName1, OriClientName2, ... → OriClientName
    OppoClientName1, OppoClientName2, ... → OppOriClientName
    """
    ori_names = []
    oppo_names = []
    
    for key in sorted(input_data.keys()):
        if key.startswith("OriClientName"):
            suffix = key[len("OriClientName"):]
            if suffix.isdigit():
                name = input_data[key]
                if name:
                    ori_names.append(name)
    
    for key in sorted(input_data.keys()):
        if key.startswith("OppoClientName"):
            suffix = key[len("OppoClientName"):]
            if suffix.isdigit():
                name = input_data[key]
                if name:
                    oppo_names.append(name)
    
    return ori_names, oppo_names


def transform_questionnaire_to_filler_data(questionnaire_answers: Dict[str, Any]) -> Dict[str, Any]:
    """
    Transform questionnaire answers into the format expected by the filler.
    Maps question IDs to template field names.
    """
    # Mapping from questionnaire question IDs to filler field names
    FIELD_MAPPING = {
        # Personal info
        "q1": "OriClientName1",
        "q3": "OriClientGender1",
        "q4": "OriClientDOB1",
        "q5": "OriClientRace1",
        "q7": "OriClientAddress1",
        "q8": "OriClientTele1",
        "q9": "OriClientWorkdet1",
        # ID card OCR data
        "id_card_name": "OriClientName1",
        "id_card_gender": "OriClientGender1",
        "id_card_dob": "OriClientDOB1",
        "id_card_race": "OriClientRace1",
        "id_card_address": "OriClientAddress1",
        "id_card_number": "OriClientIDNumber1",
        # Accident info
        "q2": "AccidentType",
        "q10": "AccidentDate",
        "q11": "AccidentLocation",
        "q12": "AccidentDescription",
        # Opponent info
        "q13": "OppoClientName1",
        "q14": "OppoClientTele1",
        "q15": "OppoClientAddress1",
        # Insurance info
        "q16": "InsuranceCompany",
        "q17": "InsurancePolicyNumber",
        # Damages
        "q18": "DamageAmount",
        "q19": "DamageDescription",
        # Court info
        "q20": "CourtName",
        "q21": "CaseNumber",
        # Form type
        "form_type": "Form"
    }
    
    output = {}
    
    for q_id, answer in questionnaire_answers.items():
        # Extract value from answer dict
        if isinstance(answer, dict):
            value = answer.get("value", answer.get("text", ""))
        else:
            value = answer
        
        # Map to filler field
        if q_id in FIELD_MAPPING:
            output[FIELD_MAPPING[q_id]] = value
        else:
            # Keep unmapped fields with original key
            output[q_id] = value
    
    return output


def transform_json(input_data: dict) -> dict:
    """
    Main transformation function.
    Converts individual fields to formatted paragraphs.
    """
    output = dict(input_data)
    
    form = input_data.get("Form", "异议")
    
    # Summarize client names
    ori_names, oppo_names = summarize_client_names(input_data)
    output["OriClientName"] = "、".join(ori_names) if ori_names else ""
    output["OppOriClientName"] = "、".join(oppo_names) if oppo_names else ""
    
    # Process OriClient
    ori_clients = extract_numbered_fields(input_data, "OriClient")
    if ori_clients:
        client_texts = []
        for num in sorted(ori_clients.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            client_texts.append(format_oriclient(ori_clients[num], form=form, is_opposing=False))
        output["OriClientInfo"] = "\n\t".join(client_texts)
    else:
        output["OriClientInfo"] = ""
    
    # Process OppoClient
    if has_oppoclient_data(input_data):
        oppo_clients = extract_numbered_fields(input_data, "OppoClient")
        if oppo_clients:
            client_texts = []
            for num in sorted(oppo_clients.keys(), key=lambda x: int(x) if x.isdigit() else 0):
                client_texts.append(format_oriclient(oppo_clients[num], form=form, is_opposing=True))
            output["OppOriClientInfo"] = "\n\t".join(client_texts)
    else:
        output["OppOriClientInfo"] = ""
    
    # Process RealClient
    real_clients = extract_numbered_fields(input_data, "RealClient")
    if real_clients:
        client_texts = []
        for num in sorted(real_clients.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            client_texts.append(format_realclient(real_clients[num]))
        output["RealClientInfo"] = "\n\t".join(client_texts)
    else:
        output["RealClientInfo"] = ""
    
    # Process Agent
    agents = extract_numbered_fields(input_data, "Agent")
    if agents:
        agent_texts = []
        for num in sorted(agents.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            agent_texts.append(format_agent(agents[num]))
        output["AgentInfo"] = "\n\t".join(agent_texts)
    else:
        output["AgentInfo"] = ""
    
    # Add date fields
    today = datetime.now()
    output.setdefault("CurrentYear", str(today.year))
    output.setdefault("CurrentMonth", str(today.month))
    output.setdefault("CurrentDay", str(today.day))
    output.setdefault("CurrentDate", today.strftime("%Y年%m月%d日"))
    
    return output


# ============================================================================
# DOCX FILLING FUNCTIONS
# ============================================================================

def find_placeholders(text: str) -> List[dict]:
    """Find all placeholder patterns in text"""
    placeholders = []
    
    # Pattern 1: {{field_name}}
    for match in re.finditer(r'\{\{(\w+)\}\}', text):
        placeholders.append({
            "start": match.start(),
            "end": match.end(),
            "field": match.group(1),
            "pattern": match.group(0)
        })
    
    # Pattern 2: {field_name}
    for match in re.finditer(r'\{(\w+)\}', text):
        # Skip if it's part of {{}}
        if text[match.start()-1:match.start()] == "{":
            continue
        placeholders.append({
            "start": match.start(),
            "end": match.end(),
            "field": match.group(1),
            "pattern": match.group(0)
        })
    
    # Pattern 3: need to fill: field_name
    for match in re.finditer(r'need to fill:\s*(\w+)', text, re.IGNORECASE):
        placeholders.append({
            "start": match.start(),
            "end": match.end(),
            "field": match.group(1),
            "pattern": match.group(0)
        })
    
    return sorted(placeholders, key=lambda x: x["start"])


def fill_text_with_data(text: str, data: dict) -> str:
    """Fill all placeholders in text with data values"""
    result = text
    
    # Replace {{field}} patterns
    for key, value in data.items():
        if value is not None:
            str_value = str(value)
            result = result.replace(f"{{{{{key}}}}}", str_value)
            result = result.replace(f"{{{key}}}", str_value)
            result = re.sub(f"need to fill:\\s*{key}", str_value, result, flags=re.IGNORECASE)
    
    return result


def apply_fangsong_to_document(doc: DocxDocument, fangsong_font: str = "仿宋_GB2312") -> Tuple[int, int, int]:
    """Apply FangSong font to document body text (except headers)"""
    changed_runs = 0
    table_runs = 0
    skipped_paras = 0
    
    for para in doc.paragraphs:
        # Skip headings
        if para.style and para.style.name.startswith("Heading"):
            skipped_paras += 1
            continue
        
        # Skip centered paragraphs (likely titles)
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            skipped_paras += 1
            continue
        
        for run in para.runs:
            run.font.name = fangsong_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fangsong_font)
            changed_runs += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = fangsong_font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), fangsong_font)
                        table_runs += 1
    
    return changed_runs, table_runs, skipped_paras


def remove_empty_lines(doc: DocxDocument):
    """Remove empty paragraphs from document"""
    paragraphs_to_remove = []
    
    for para in doc.paragraphs:
        if not para.text.strip():
            # Check if it's not a page break or section break
            if not any(child.tag.endswith('br') for child in para._element):
                paragraphs_to_remove.append(para)
    
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)


def fill_template(
    template_path: str,
    json_data: dict,
    output_path: str,
    force_font: Optional[str] = None,
    apply_fangsong: bool = True,
    fangsong_font: str = "仿宋_GB2312"
) -> Tuple[bool, int]:
    """Fill the template with JSON data while preserving formatting."""
    doc = DocxDocument(template_path)
    filled_count = 0
    
    # Process all paragraphs
    for para in doc.paragraphs:
        original_text = para.text
        if not original_text:
            continue
        
        filled_text = fill_text_with_data(original_text, json_data)
        
        if filled_text != original_text:
            # Clear and rebuild paragraph
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = filled_text
            else:
                para.add_run(filled_text)
            filled_count += 1
    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    if not original_text:
                        continue
                    
                    filled_text = fill_text_with_data(original_text, json_data)
                    
                    if filled_text != original_text:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = filled_text
                        else:
                            para.add_run(filled_text)
                        filled_count += 1
    
    # Apply FangSong font if enabled
    if apply_fangsong:
        apply_fangsong_to_document(doc, fangsong_font)
    
    # Remove empty lines
    remove_empty_lines(doc)
    
    # Save filled document
    doc.save(output_path)
    return True, filled_count


# ============================================================================
# Document Filler Service Class
# ============================================================================

class DocumentFillerService:
    """Service for filling document templates with questionnaire data"""
    
    def __init__(self, templates_dir: str = None, output_dir: str = None):
        self.templates_dir = Path(templates_dir or "templates")
        self.output_dir = Path(output_dir or "uploads/generated_documents")
        
        # Create directories if they don't exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def get_template_path(self, template_code: str) -> Optional[Path]:
        """Get the path to a template file by code"""
        # Try different naming conventions
        patterns = [
            f"{template_code}.docx",
            f"{template_code}_*.docx",
            f"*_{template_code}.docx"
        ]
        
        for pattern in patterns:
            matches = list(self.templates_dir.glob(pattern))
            if matches:
                return matches[0]
        
        return None
    
    async def fill_from_questionnaire(
        self,
        template_code: str,
        questionnaire_answers: Dict[str, Any],
        autofill_data: Optional[Dict[str, Any]] = None,
        apply_fangsong: bool = True,
        session_id: Optional[str] = None,
        user_id: Optional[str] = None,
        case_uuid: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Fill a template using questionnaire answers.
        
        Args:
            template_code: Code identifying the template
            questionnaire_answers: Answers from the questionnaire
            autofill_data: Additional data from OCR (e.g., ID card)
            apply_fangsong: Whether to apply FangSong font
            session_id: Optional session ID for linking
            user_id: Optional user ID for ownership
            case_uuid: Optional case UUID for linking
        
        Returns:
            Dict with success, document_id, output_filename, download_url, error
        """
        # Find template
        template_path = self.get_template_path(template_code)
        if not template_path:
            return {
                "success": False,
                "error": f"模板未找到: {template_code}"
            }
        
        # Extract actual values from questionnaire answers
        # Answers are in format {question_id: {value: "xxx", answered_at: "..."}}
        clean_answers = {}
        for q_id, ans in questionnaire_answers.items():
            if isinstance(ans, dict):
                clean_answers[q_id] = ans.get("value", "")
            else:
                clean_answers[q_id] = ans
        
        # Transform questionnaire answers to filler format
        filler_data = transform_questionnaire_to_filler_data(clean_answers)
        
        # Merge with autofill data (OCR results take priority for matching fields)
        if autofill_data:
            # Transform OCR data to filler format
            ocr_filler_data = transform_questionnaire_to_filler_data(autofill_data)
            filler_data.update(ocr_filler_data)
        
        # Apply full transformation (handles OriClient formatting, etc.)
        transformed_data = transform_json(filler_data)
        
        # Generate output filename
        client_name = transformed_data.get("OriClientName1", "")
        # Clean filename
        client_name = re.sub(r'[\\/:*?"<>|]', '', client_name)[:20]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_uuid = str(uuid.uuid4())[:8]
        
        if client_name:
            output_filename = f"{template_code}_{client_name}_{timestamp}_{doc_uuid}.docx"
        else:
            output_filename = f"{template_code}_{timestamp}_{doc_uuid}.docx"
        
        output_path = self.output_dir / output_filename
        
        # Fill the template
        try:
            success, filled_count = fill_template(
                str(template_path),
                transformed_data,
                str(output_path),
                apply_fangsong=apply_fangsong
            )
            
            if success:
                # Generate document ID
                document_id = str(uuid.uuid4())
                
                # Generate download URL (relative path for API)
                download_url = f"/api/documents/download/{document_id}"
                
                return {
                    "success": True,
                    "document_id": document_id,
                    "output_filename": output_filename,
                    "output_path": str(output_path),
                    "download_url": download_url,
                    "filled_fields": filled_count,
                    "template_code": template_code,
                    "session_id": session_id,
                    "user_id": user_id,
                    "case_uuid": case_uuid
                }
            else:
                return {
                    "success": False,
                    "error": "文档填充失败"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"生成文档时出错: {str(e)}"
            }
    
    def fill_directly(
        self,
        template_path: str,
        data: Dict[str, Any],
        output_path: Optional[str] = None,
        apply_fangsong: bool = True
    ) -> FillerResult:
        """
        Fill a template directly with provided data.
        
        Args:
            template_path: Path to the template file
            data: Data to fill (already in correct format)
            output_path: Optional output path
            apply_fangsong: Whether to apply FangSong font
        """
        if not os.path.exists(template_path):
            return FillerResult(
                success=False,
                error=f"Template file not found: {template_path}"
            )
        
        # Generate output path if not provided
        if not output_path:
            template_name = Path(template_path).stem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{template_name}_filled_{timestamp}.docx"
            output_path = str(self.output_dir / output_filename)
        
        # Transform and fill
        transformed_data = transform_json(data)
        
        try:
            success, filled_count = fill_template(
                template_path,
                transformed_data,
                output_path,
                apply_fangsong=apply_fangsong
            )
            
            return FillerResult(
                success=success,
                output_path=output_path,
                output_filename=Path(output_path).name,
                filled_fields=filled_count
            )
        except Exception as e:
            return FillerResult(
                success=False,
                error=str(e)
            )


# ============================================================================
# Singleton instance
# ============================================================================

_filler_service: Optional[DocumentFillerService] = None


def get_filler_service(templates_dir: str = None, output_dir: str = None) -> DocumentFillerService:
    """Get or create the document filler service instance"""
    global _filler_service
    if _filler_service is None:
        _filler_service = DocumentFillerService(templates_dir, output_dir)
    return _filler_service
